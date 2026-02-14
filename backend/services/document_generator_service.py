from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from jinja2 import Template, Environment, FileSystemLoader
from io import BytesIO
from typing import List, Dict, Any
from datetime import datetime
from pathlib import Path

# Dodaj importy dla ReportLab
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_CENTER, TA_RIGHT

class DocumentGeneratorService:
    """
    Serwis do generowania dokumentów kart wyjazdów w różnych formatach
    """
    
    def __init__(self):
        """Inicjalizacja - rejestruj czcionki dla polskich znaków"""
        try:
            # Spróbuj zarejestrować czcionkę Arial (Windows)
            pdfmetrics.registerFont(TTFont('Arial', 'arial.ttf'))
            pdfmetrics.registerFont(TTFont('Arial-Bold', 'arialbd.ttf'))
            self.font_name = 'Arial'
        except:
            # Fallback - użyj Helvetica (wbudowana)
            self.font_name = 'Helvetica'
        # Konfiguracja Jinja2 - ścieżka do szablonów
        template_dir = Path(__file__).parent.parent / 'templates'
        self.jinja_env = Environment(loader=FileSystemLoader(str(template_dir)))

    def generate_docx(self, firefighter_name: str, records: List[Dict[str, Any]], 
                  date_from: str = None, date_to: str = None,
                  firefighter_data: Dict[str, str] = None) -> BytesIO:
        """
        Generuje kartę wyjazdów w formacie DOCX
        ORIENTACJA POZIOMA (LANDSCAPE) - zgodnie z szablonem HTML
        """
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        doc = Document()
        
        # Ustaw orientację poziomą i marginesy
        section = doc.sections[0]
        section.page_height = Cm(21)
        section.page_width = Cm(29.7)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        
        # Przygotuj dane strażaka
        stopien = firefighter_data.get('stopien', '.....................') if firefighter_data else '.....................'
        nazwisko_imie = firefighter_data.get('nazwisko_imie', firefighter_name) if firefighter_data else firefighter_name
        stanowisko = firefighter_data.get('stanowisko', '.....................') if firefighter_data else '.....................'
        
        # Przygotuj rekordy (tylko zaliczone do emerytury)
        idx = 1
        all_records = []
        for record in records:
            if record.get('zaliczono_do_emerytury', '') != '1':
                continue
            
            czas_rozp = record.get('czas_rozp_zdarzenia', '')
            data = czas_rozp[:10] if len(czas_rozp) >= 10 else czas_rozp
            
            # Rodzaj zagrożenia
            rodzaj_zagrozenia = ''
            p_value = str(record.get('p', '')).strip()
            mz_value = str(record.get('mz', '')).strip()
            
            if p_value == '1':
                rodzaj_zagrozenia = 'pożar'
            elif mz_value == '1':
                rodzaj_zagrozenia = 'm.zagrożenie'
            
            # Zadanie podstawowe
            funkcja = record.get('funkcja', '').lower()
            zadanie_podstawowe = 'X' if 'ratownik' in funkcja else ''
            
            # Zadanie specjalistyczne (puste)
            zadanie_specjalistyczne = ''
            
            # Kierowanie
            kierowanie = 'X' if 'dowódca' in funkcja or 'dowodca' in funkcja else ''
            
            all_records.append({
                'lp': idx,
                'data': data,
                'rodzaj_zagrozenia': rodzaj_zagrozenia,
                'zadanie_podstawowe': zadanie_podstawowe,
                'zadanie_specjalistyczne': zadanie_specjalistyczne,
                'kierowanie': kierowanie,
                'nr_meldunku': record.get('nr_meldunku', ''),
            })
            idx += 1
        
        # PAGINACJA: Strona 1 = 16 wierszy, kolejne = 20 wierszy
        pages = []
        page_number = 1
        records_remaining = all_records[:]
        
        while records_remaining:
            if page_number == 1:
                page_records = records_remaining[:16]
                records_remaining = records_remaining[16:]
            else:
                page_records = records_remaining[:20]
                records_remaining = records_remaining[20:]
            
            pages.append({
                'page_number': page_number,
                'records': page_records,
                'is_last_page': len(records_remaining) == 0
            })
            page_number += 1
        
        # Generuj każdą stronę
        for page_idx, page in enumerate(pages):
            if page['page_number'] == 1:
                # STRONA 1 - Pełny nagłówek
                
                # 1. Pieczątka i numer strony
                header_table = doc.add_table(rows=2, cols=2)
                header_table.autofit = False
                header_table.columns[0].width = Cm(5)
                header_table.columns[1].width = Cm(20)
                
                # Lewa kolumna - pieczątka
                cell_stamp = header_table.cell(0, 0)
                cell_stamp.text = ""
                # Bottom border dotted
                tc = cell_stamp._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'dotted')   # styl wykropkowany
                bottom.set(qn('w:sz'), '8')         # grubość (np. 8 = 1pt)
                bottom.set(qn('w:space'), '0')
                bottom.set(qn('w:color'), 'auto')
                tcBorders.append(bottom)
                tcPr.append(tcBorders)
                
                cell_stamp_label = header_table.cell(1, 0)
                p = cell_stamp_label.paragraphs[0]
                run = p.add_run("(pieczątka jednostki organizacyjnej)")
                run.font.size = Pt(8)
                run.font.italic = True
                run.font.color.rgb = RGBColor(128, 128, 128)
                
                # Prawa kolumna - numer strony
                cell_page = header_table.cell(0, 1)
                p = cell_page.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = p.add_run(f"Strona nr {page['page_number']}")
                run.font.size = Pt(8)
                
                # Usuń ramki z header_table
                self._remove_table_borders(header_table)
                
                doc.add_paragraph()
                
                # 2. Tytuł
                title1 = doc.add_paragraph()
                title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = title1.add_run("Roczna karta ewidencji bezpośredniego udziału strażaka w działaniach ratowniczych,")
                run.font.size = Pt(11)
                run.font.bold = True
                
                title2 = doc.add_paragraph()
                title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = title2.add_run("w tym ratowniczo-gaśniczych lub bezpośredniego kierowania tymi działaniami na miejscu zdarzenia")
                run.font.size = Pt(11)
                run.font.bold = True
                
                # 3. Zakres dat
                date_para = doc.add_paragraph()
                date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                date_text = f"od {date_from if date_from else '.....................'} do {date_to if date_to else '.....................'}"
                run = date_para.add_run(date_text)
                run.font.size = Pt(10)
                
                # 4. Dane strażaka
                ff_para = doc.add_paragraph()
                ff_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = ff_para.add_run(f"{stopien} {nazwisko_imie} - {stanowisko}")
                run.font.size = Pt(10)
                
                ff_label = doc.add_paragraph()
                ff_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = ff_label.add_run("(stopień, tytuł, imię, nazwisko, stanowisko)")
                run.font.size = Pt(8)
                run.font.italic = True
                run.font.color.rgb = RGBColor(128, 128, 128)
                
                doc.add_paragraph()
            
            else:
                # STRONY 2+ - Uproszczony nagłówek
                if page_idx > 0:
                    # Dodaj page break przed nową stroną
                    doc.add_page_break()
                
                header_table = doc.add_table(rows=2, cols=2)
                header_table.autofit = False
                header_table.columns[0].width = Cm(20)
                header_table.columns[1].width = Cm(5)
                
                # Lewa kolumna - dane strażaka
                cell_ff = header_table.cell(0, 0)
                p = cell_ff.paragraphs[0]
                run = p.add_run(f"{stopien} {nazwisko_imie} - {stanowisko}")
                run.font.size = Pt(10)
                
                cell_ff_label = header_table.cell(1, 0)
                p = cell_ff_label.paragraphs[0]
                run = p.add_run("(stopień, tytuł, imię, nazwisko, stanowisko)")
                run.font.size = Pt(8)
                run.font.italic = True
                run.font.color.rgb = RGBColor(128, 128, 128)
                
                # Prawa kolumna - numer strony
                cell_page = header_table.cell(0, 1)
                p = cell_page.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = p.add_run(f"Strona nr {page['page_number']}")
                run.font.size = Pt(9)
                
                # Usuń ramki
                self._remove_table_borders(header_table)
                
                doc.add_paragraph()
            
            # TABELA GŁÓWNA
            # Nagłówek: 2 wiersze
            table = doc.add_table(rows=2, cols=8)
            table.style = 'Table Grid'
            
            # Wiersz 1 nagłówka - z merged cells
            hdr_cells = table.rows[0].cells
            
            # Lp. - rowspan 2
            hdr_cells[0].text = 'Lp.'
            hdr_cells[0].merge(table.rows[1].cells[0])
            
            # Data - rowspan 2
            hdr_cells[1].text = 'Data'
            hdr_cells[1].merge(table.rows[1].cells[1])
            
            # Rodzaj zagrożenia - rowspan 2
            hdr_cells[2].text = 'Rodzaj zagrożenia\n(pożar, m. zagrożenie)'
            hdr_cells[2].merge(table.rows[1].cells[2])
            
            # Wykonywanie zadań - colspan 3
            hdr_cells[3].text = 'Wykonywanie zadań*'
            hdr_cells[3].merge(hdr_cells[4])
            hdr_cells[3].merge(hdr_cells[5])
            
            # Nazwa dokumentu - rowspan 2
            hdr_cells[6].text = 'Nazwa dokumentu źródłowego'
            hdr_cells[6].merge(table.rows[1].cells[6])
            
            # Podpis - rowspan 2
            hdr_cells[7].text = 'Podpis strażaka biorącego udział w działaniach ratowniczych, w tym ratowniczo-gaśniczych lub bezpośrednio kierującego tymi działaniami na miejscu zdarzenia'
            hdr_cells[7].merge(table.rows[1].cells[7])
            
            # Wiersz 2 nagłówka - subheadery
            hdr_row2 = table.rows[1].cells
            hdr_row2[3].text = 'Podstawowych'
            hdr_row2[4].text = 'Specjalistycznych**'
            hdr_row2[5].text = 'Kierowanie działaniami ratowniczymi'
            hdr_row2[6].text = '(nr meldunku - inf. ze zdarzenia)'
            
            # Styl nagłówków
            for row_idx in range(2):
                for cell in table.rows[row_idx].cells:
                    self._set_cell_background(cell, "D3D3D3")
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(7)
            
            # Dane
            for record in page['records']:
                row_cells = table.add_row().cells
                
                data_row = [
                    str(record['lp']),
                    record['data'],
                    record['rodzaj_zagrozenia'],
                    record['zadanie_podstawowe'],
                    record['zadanie_specjalistyczne'],
                    record['kierowanie'],
                    record['nr_meldunku'],
                    ''
                ]
                
                for i, value in enumerate(data_row):
                    row_cells[i].text = value
                    for paragraph in row_cells[i].paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.size = Pt(7)
            
            # Szerokości kolumn
            table.columns[0].width = Cm(0.6)   # Lp
            table.columns[1].width = Cm(2.5)   # Data
            table.columns[2].width = Cm(3)     # Rodzaj zagrożenia
            table.columns[3].width = Cm(2.5)   # Podstawowych
            table.columns[4].width = Cm(2.8)   # Specjalistycznych
            table.columns[5].width = Cm(2.8)   # Kierowanie
            table.columns[6].width = Cm(3)     # Nr meldunku
            table.columns[7].width = Cm(7.8)   # Podpis
            
            doc.add_paragraph()
            
            # STOPKA TABELI
            footer_table = doc.add_table(rows=3, cols=5)
            footer_table.autofit = False
            
            # Kolumny: data | margin | sprawdził | margin | zatwierdził
            footer_table.columns[0].width = Cm(7.5)
            footer_table.columns[1].width = Cm(1.25)
            footer_table.columns[2].width = Cm(8.75)
            footer_table.columns[3].width = Cm(1.25)
            footer_table.columns[4].width = Cm(6.25)
            
            # Wiersz 1 - etykiety
            footer_table.cell(0, 0).text = ""
            p = footer_table.cell(0, 2).paragraphs[0]
            run = p.add_run("Sprawdził:")
            run.font.size = Pt(7)
            
            p = footer_table.cell(0, 4).paragraphs[0]
            run = p.add_run("Zatwierdził:")
            run.font.size = Pt(7)
            
            # Wiersz 2 - miejsce na podpis (puste z dużym paddingiem)
            footer_table.cell(1, 0).text = ""
            footer_table.cell(1, 2).text = ""
            footer_table.cell(1, 4).text = ""
            
            # Dodaj dolną kreskę do wiersza 2
            for col in [0, 2, 4]:
                cell = footer_table.cell(1, col)
                tcPr = cell._element.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '6')
                bottom.set(qn('w:color'), '808080')
                tcBorders.append(bottom)
                tcPr.append(tcBorders)
            
            # Wiersz 3 - opisy
            p = footer_table.cell(2, 0).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("(miejscowość, data)")
            run.font.size = Pt(7)
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
            
            p = footer_table.cell(2, 2).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("(podpis i pieczątka kierownika komórki operacyjnej\nlub dowódcy jednostki ratowniczo-gaśniczej)")
            run.font.size = Pt(7)
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
            
            p = footer_table.cell(2, 4).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("(podpis i pieczątka kierownika jednostki organizacyjnej)")
            run.font.size = Pt(7)
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
            
            # Usuń ramki z footer_table
            self._remove_table_borders(footer_table)
            
            # Objaśnienia
            doc.add_paragraph()
            
            p = doc.add_paragraph()
            run = p.add_run("Objaśnienia:")
            run.font.size = Pt(7)
            run.font.bold = True
            
            p = doc.add_paragraph()
            run = p.add_run("* - wpisać znak X")
            run.font.size = Pt(7)
            
            p = doc.add_paragraph()
            run = p.add_run("** - dotyczy udziału w działaniach ratowniczych w ramach Specjalistycznych grup roboczych")
            run.font.size = Pt(7)
        
        # Zapisz
        output = BytesIO()
        doc.save(output)
        output.seek(0)
    
        return output

    def _remove_table_borders(self, table):
        """Usuń ramki z tabeli"""
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tblBorders.append(border)
        
        tblPr.append(tblBorders)
    
    def _set_cell_background(self, cell, color):
        """Ustaw kolor tła komórki"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    def generate_html(self, firefighter_name: str, records: List[Dict[str, Any]],
                    date_from: str = None, date_to: str = None,
                    firefighter_data: Dict[str, str] = None) -> str:
        """
        Generuje kartę wyjazdów w formacie HTML używając szablonu z pliku
        ORIENTACJA POZIOMA (LANDSCAPE)
        Z PAGINACJĄ: strona 1 = 16 wierszy, kolejne strony = 20 wierszy
        """
        # Załaduj szablon z pliku
        template = self.jinja_env.get_template('karta_wyjazdow.html')
        
        # Przygotuj dane strażaka
        stopien = firefighter_data.get('stopien', '.....................') if firefighter_data else '.....................'
        nazwisko_imie = firefighter_data.get('nazwisko_imie', firefighter_name) if firefighter_data else firefighter_name
        stanowisko = firefighter_data.get('stanowisko', '.....................') if firefighter_data else '.....................'
        idx = 1
        # Przygotuj dane dla szablonu
        all_records = []
        for id, record in enumerate(records, start=1):
            # Pomiń rekord jeżele nie zaliczony do emerytury
            if record.get('zaliczono_do_emerytury', '') != '1':
                continue
            # Wyciągnij datę (pierwsze 10 znaków)
            czas_rozp = record.get('czas_rozp_zdarzenia', '')
            data = czas_rozp[:10] if len(czas_rozp) >= 10 else czas_rozp
            
            # Określ rodzaj zagrożenia
            rodzaj_zagrozenia = ''
            p_value = str(record.get('p', '')).strip()
            mz_value = str(record.get('mz', '')).strip()
            
            if p_value == '1':
                rodzaj_zagrozenia = 'pożar'
            elif mz_value == '1':
                rodzaj_zagrozenia = 'm.zagrożenie'
            
            # Określ zadanie podstawowe (X jeśli ratownik)
            funkcja = record.get('funkcja', '').lower()
            zadanie_podstawowe = 'X' if 'ratownik' in funkcja else ''
            
            # Zadanie specjalistyczne - na razie puste (do implementacji po zmianie modelu)
            zadanie_specjalistyczne = ''
            
            # Kierowanie działaniami (X jeśli dowódca)
            kierowanie = 'X' if 'dowódca' in funkcja or 'dowodca' in funkcja else ''
            
            all_records.append({
                'lp': idx,  # Globalne Lp przez wszystkie strony
                'data': data,
                'rodzaj_zagrozenia': rodzaj_zagrozenia,
                'zadanie_podstawowe': zadanie_podstawowe,
                'zadanie_specjalistyczne': zadanie_specjalistyczne,
                'kierowanie': kierowanie,
                'nr_meldunku': record.get('nr_meldunku', ''),
            })
            # Zwieksz licznik
            idx += 1
        
        # PAGINACJA: Strona 1 = 16 wierszy, kolejne = 20 wierszy
        pages = []
        page_number = 1
        records_remaining = all_records[:]
        
        while records_remaining:
            if page_number == 1:
                # Pierwsza strona - 16 wierszy
                page_records = records_remaining[:16]
                records_remaining = records_remaining[16:]
            else:
                # Kolejne strony - 20 wierszy
                page_records = records_remaining[:20]
                records_remaining = records_remaining[20:]
            
            pages.append({
                'page_number': page_number,
                'records': page_records,
                'is_last_page': len(records_remaining) == 0
            })
            
            page_number += 1
        
        # Renderuj szablon
        html_content = template.render(
            firefighter_name=firefighter_name,
            stopien=stopien,
            nazwisko_imie=nazwisko_imie,
            stanowisko=stanowisko,
            date_from=date_from if date_from else '.....................',
            date_to=date_to if date_to else '.....................',
            pages=pages
        )
        
        return html_content
    
    def generate_pdf(self, firefighter_name: str, records: List[Dict[str, Any]], 
                 date_from: str = None, date_to: str = None, 
                 firefighter_data: Dict[str, str] = None) -> BytesIO:
        """
        Generuje kartę wyjazdów w formacie PDF używając ReportLab
        ORIENTACJA POZIOMA (LANDSCAPE) - zgodnie z szablonem HTML
        """
        from reportlab.lib.pagesizes import landscape, A4
        from reportlab.platypus import PageBreak
        
        output = BytesIO()
        
        # Utwórz dokument PDF w orientacji poziomej
        doc = SimpleDocTemplate(
            output,
            pagesize=landscape(A4),
            topMargin=1.5*cm,
            bottomMargin=1.5*cm,
            leftMargin=1.5*cm,
            rightMargin=1.5*cm
        )
        
        # Elementy dokumentu
        elements = []
        
        # Style
        styles = getSampleStyleSheet()
        
        # Przygotuj dane strażaka
        stopien = firefighter_data.get('stopien', '.....................') if firefighter_data else '.....................'
        nazwisko_imie = firefighter_data.get('nazwisko_imie', firefighter_name) if firefighter_data else firefighter_name
        stanowisko = firefighter_data.get('stanowisko', '.....................') if firefighter_data else '.....................'
        
        # Przygotuj rekordy (tylko zaliczone do emerytury)
        idx = 1
        all_records = []
        for record in records:
            if record.get('zaliczono_do_emerytury', '') != '1':
                continue
            
            czas_rozp = record.get('czas_rozp_zdarzenia', '')
            data = czas_rozp[:10] if len(czas_rozp) >= 10 else czas_rozp
            
            # Rodzaj zagrożenia
            rodzaj_zagrozenia = ''
            p_value = str(record.get('p', '')).strip()
            mz_value = str(record.get('mz', '')).strip()
            
            if p_value == '1':
                rodzaj_zagrozenia = 'pożar'
            elif mz_value == '1':
                rodzaj_zagrozenia = 'm.zagrożenie'
            
            # Zadanie podstawowe
            funkcja = record.get('funkcja', '').lower()
            zadanie_podstawowe = 'X' if 'ratownik' in funkcja else ''
            
            # Zadanie specjalistyczne (puste)
            zadanie_specjalistyczne = ''
            
            # Kierowanie
            kierowanie = 'X' if 'dowódca' in funkcja or 'dowodca' in funkcja else ''
            
            all_records.append({
                'lp': idx,
                'data': data,
                'rodzaj_zagrozenia': rodzaj_zagrozenia,
                'zadanie_podstawowe': zadanie_podstawowe,
                'zadanie_specjalistyczne': zadanie_specjalistyczne,
                'kierowanie': kierowanie,
                'nr_meldunku': record.get('nr_meldunku', ''),
            })
            idx += 1
        
        # PAGINACJA: Strona 1 = 16 wierszy, kolejne = 20 wierszy
        pages = []
        page_number = 1
        records_remaining = all_records[:]
        
        while records_remaining:
            if page_number == 1:
                page_records = records_remaining[:16]
                records_remaining = records_remaining[16:]
            else:
                page_records = records_remaining[:20]
                records_remaining = records_remaining[20:]
            
            pages.append({
                'page_number': page_number,
                'records': page_records,
                'is_last_page': len(records_remaining) == 0
            })
            page_number += 1
        
        # Generuj każdą stronę
        for page in pages:
            if page['page_number'] == 1:
                # STRONA 1 - Pełny nagłówek
                # 1. Pieczątka i numer strony
                header_data = [
                    [
                        Paragraph("........................................", styles['Normal']),
                        Paragraph(f"Strona nr {page['page_number']}", ParagraphStyle(
                            'PageNum',
                            parent=styles['Normal'],
                            fontSize=9,
                            alignment=TA_RIGHT
                        ))
                    ],
                    [
                        Paragraph("<i>(pieczątka jednostki organizacyjnej)</i>", ParagraphStyle(
                            'Stamp',
                            parent=styles['Normal'],
                            fontSize=8,
                            textColor=colors.grey
                        )),
                        ""
                    ]
                ]
                
                header_table = Table(header_data, colWidths=[15*cm, 10*cm])
                header_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                    ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                elements.append(header_table)
                elements.append(Spacer(1, 0.3*cm))
                
                # 2. Tytuł
                title_style = ParagraphStyle(
                    'Title',
                    parent=styles['Normal'],
                    fontSize=11,
                    alignment=TA_CENTER,
                    fontName='Helvetica-Bold',
                    leading=14
                )
                
                title = Paragraph(
                    "Roczna karta ewidencji bezpośredniego udziału strażaka w działaniach ratowniczych,<br/>"
                    "w tym ratowniczo-gaśniczych lub bezpośredniego kierowania tymi działaniami na miejscu zdarzenia",
                    title_style
                )
                elements.append(title)
                elements.append(Spacer(1, 0.3*cm))
                
                # 3. Zakres dat
                date_style = ParagraphStyle(
                    'DateRange',
                    parent=styles['Normal'],
                    fontSize=10,
                    alignment=TA_CENTER
                )
                date_text = f"od {date_from if date_from else '.....................'} do {date_to if date_to else '.....................'}"
                elements.append(Paragraph(date_text, date_style))
                elements.append(Spacer(1, 0.2*cm))
                
                # 4. Dane strażaka
                firefighter_style = ParagraphStyle(
                    'Firefighter',
                    parent=styles['Normal'],
                    fontSize=10,
                    alignment=TA_CENTER
                )
                firefighter_label_style = ParagraphStyle(
                    'FirefighterLabel',
                    parent=styles['Normal'],
                    fontSize=8,
                    textColor=colors.grey,
                    alignment=TA_CENTER
                )
                
                firefighter_text = f"{stopien} {nazwisko_imie} - {stanowisko}"
                elements.append(Paragraph(firefighter_text, firefighter_style))
                elements.append(Paragraph("<i>(stopień, tytuł, imię, nazwisko, stanowisko)</i>", firefighter_label_style))
                elements.append(Spacer(1, 0.3*cm))
            
            else:
                # STRONY 2+ - Uproszczony nagłówek
                header_left_style = ParagraphStyle(
                    'HeaderLeft',
                    parent=styles['Normal'],
                    fontSize=10
                )
                header_label_style = ParagraphStyle(
                    'HeaderLabel',
                    parent=styles['Normal'],
                    fontSize=8,
                    textColor=colors.grey
                )
                header_right_style = ParagraphStyle(
                    'HeaderRight',
                    parent=styles['Normal'],
                    fontSize=9,
                    alignment=TA_RIGHT
                )
                
                header_data = [
                    [
                        Paragraph(f"{stopien} {nazwisko_imie} - {stanowisko}", header_left_style),
                        Paragraph(f"Strona nr {page['page_number']}", header_right_style)
                    ],
                    [
                        Paragraph("<i>(stopień, tytuł, imię, nazwisko, stanowisko)</i>", header_label_style),
                        ""
                    ]
                ]
                
                header_table = Table(header_data, colWidths=[20*cm, 5*cm])
                header_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                    ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                elements.append(header_table)
                elements.append(Spacer(1, 0.3*cm))
            
            # TABELA GŁÓWNA
            table_data = [
                # Nagłówek 1
                ['Lp.', 'Data', 'Rodzaj zagrożenia\n(pożar, m. zagrożenie)', 
                'Podstawowych', 'Specjalistycznych**', 'Kierowanie działaniami\nratowniczymi',
                'Nazwa dokumentu źródłowego\n(nr meldunku - inf. ze zdarzenia)',
                'Podpis strażaka biorącego udział w działaniach ratowniczych,\nw tym ratowniczo-gaśniczych lub bezpośrednio\nkierującego tymi działaniami na miejscu zdarzenia']
            ]
            
            # Dane
            for record in page['records']:
                table_data.append([
                    str(record['lp']),
                    record['data'],
                    record['rodzaj_zagrozenia'],
                    record['zadanie_podstawowe'],
                    record['zadanie_specjalistyczne'],
                    record['kierowanie'],
                    record['nr_meldunku'],
                    ''
                ])
            
            # Szerokości kolumn (dopasowane do landscape A4)
            col_widths = [0.6*cm, 2.5*cm, 3*cm, 2.5*cm, 2.8*cm, 2.8*cm, 3*cm, 7.8*cm]
            
            table = Table(table_data, colWidths=col_widths)
            table.setStyle(TableStyle([
                # Nagłówek
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 7),
                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),
                
                # Dane
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 7),
                ('ALIGN', (0, 1), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 1), (-1, -1), 'MIDDLE'),
                
                # Ramki
                ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                
                # Padding
                ('LEFTPADDING', (0, 0), (-1, -1), 2),
                ('RIGHTPADDING', (0, 0), (-1, -1), 2),
                ('TOPPADDING', (0, 0), (-1, -1), 2),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
            ]))
            
            elements.append(table)
            elements.append(Spacer(1, 0.5*cm))
            
            # STOPKA TABELI
            footer_data = [
                ['', 'Sprawdził:', 'Zatwierdził:'],
                ['', '', ''],
                ['(miejscowość, data)', 
                '(podpis i pieczątka kierownika komórki operacyjnej\nlub dowódcy jednostki ratowniczo-gaśniczej)',
                '(podpis i pieczątka kierownika jednostki organizacyjnej)']
            ]
            
            footer_table = Table(footer_data, colWidths=[7.5*cm, 8.75*cm, 8.75*cm])
            footer_table.setStyle(TableStyle([
                ('FONTSIZE', (0, 0), (-1, -1), 7),
                ('ALIGN', (0, 0), (-1, 0), 'LEFT'),
                ('ALIGN', (0, 1), (-1, 1), 'CENTER'),
                ('ALIGN', (0, 2), (-1, 2), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('BOTTOMPADDING', (0, 1), (-1, 1), 40),
                ('LINEBELOW', (0, 1), (-1, 1), 0.5, colors.grey),
                ('TEXTCOLOR', (0, 2), (-1, 2), colors.grey),
            ]))
            
            elements.append(footer_table)
            elements.append(Spacer(1, 0.3*cm))
            
            # Objaśnienia
            explanation_style = ParagraphStyle(
                'Explanation',
                parent=styles['Normal'],
                fontSize=7
            )
            elements.append(Paragraph("<b>Objaśnienia:</b>", explanation_style))
            elements.append(Paragraph("* - wpisać znak X", explanation_style))
            elements.append(Paragraph("** - dotyczy udziału w działaniach ratowniczych w ramach Specjalistycznych grup roboczych", explanation_style))
            
            # Page break jeśli nie ostatnia strona
            if not page['is_last_page']:
                elements.append(PageBreak())
        
        # Zbuduj PDF
        doc.build(elements)
        output.seek(0)
        
        return output